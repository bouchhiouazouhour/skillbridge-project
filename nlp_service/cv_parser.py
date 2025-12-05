import re
import io
import tempfile
import os
from typing import Dict, List, Any, Union
import PyPDF2
import docx
import pdfplumber

class CVParser:
    """Parse PDF/DOCX files and extract structured data"""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse CV file and extract structured data"""
        text = self._extract_text(file_path)
        
        return {
            'full_text': text,
            'contact': self._extract_contact(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
        }
    
    def extract_text(self, file_obj) -> str:
        """
        Extract text from a Flask file object (PDF or DOCX).
        Handles file uploads directly from request without saving to disk permanently.
        
        Args:
            file_obj: Flask FileStorage object from request.files
            
        Returns:
            str: Extracted text from the document
            
        Raises:
            ValueError: If file format is unsupported or file is corrupted/unreadable
        """
        if file_obj is None:
            raise ValueError("No file provided")
        
        filename = (file_obj.filename or '').lower()
        
        if not filename:
            raise ValueError("File has no filename")
        
        # Read file content into memory
        try:
            file_content = file_obj.read()
            file_obj.seek(0)  # Reset file pointer for potential re-read
        except Exception as e:
            raise ValueError(f"Failed to read file: {str(e)}")
        
        if len(file_content) == 0:
            raise ValueError("File is empty")
        
        # Handle based on file extension
        if filename.endswith('.pdf'):
            return self._extract_pdf_from_bytes(file_content)
        elif filename.endswith(('.doc', '.docx')):
            return self._extract_docx_from_bytes(file_content)
        else:
            raise ValueError(f"Unsupported file format. Please upload PDF or DOCX files. Got: {filename}")
    
    def _extract_pdf_from_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF bytes using pdfplumber with PyPDF2 fallback"""
        text = ""
        
        # Try pdfplumber first
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            # Fallback to PyPDF2
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                raise ValueError(f"Failed to extract text from PDF. The file may be corrupted or password-protected: {str(e)}")
        
        return text.strip()
    
    def _extract_docx_from_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX. The file may be corrupted: {str(e)}")
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX file (legacy method for file paths)"""
        if file_path.endswith('.pdf'):
            return self._extract_pdf_text(file_path)
        elif file_path.endswith(('.doc', '.docx')):
            return self._extract_docx_text(file_path)
        else:
            raise ValueError("Unsupported file format")
    
    def extract_text_from_path(self, file_path: str) -> str:
        """
        Public method to extract text from a file path.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            str: Extracted text from the document
        """
        return self._extract_text(file_path)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except (IOError, ValueError, KeyError) as e:
            # Fallback to PyPDF2 for specific PDF parsing errors
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() or ""
            except Exception as fallback_error:
                raise ValueError(f"Failed to extract text from PDF: {str(fallback_error)}")
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    
    def _extract_contact(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        contact = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Extract phone
        phone_pattern = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact['phone'] = phones[0]
        
        # Extract name (first line typically)
        lines = text.split('\n')
        for line in lines[:5]:
            if line.strip() and len(line.strip()) < 50:
                contact['name'] = line.strip()
                break
        
        return contact
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience"""
        experience = []
        
        # Look for experience section
        exp_section = self._find_section(text, ['experience', 'work history', 'employment'])
        
        if exp_section:
            # Simple extraction: split by bullet points or line breaks
            entries = re.split(r'\n(?=[•\-\*]|\d+\.)', exp_section)
            for entry in entries:
                if len(entry.strip()) > 20:
                    experience.append({
                        'description': entry.strip()
                    })
        
        return experience
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education"""
        education = []
        
        # Look for education section
        edu_section = self._find_section(text, ['education', 'academic', 'qualifications'])
        
        if edu_section:
            entries = re.split(r'\n(?=[•\-\*]|\d+\.)', edu_section)
            for entry in entries:
                if len(entry.strip()) > 10:
                    education.append({
                        'description': entry.strip()
                    })
        
        return education
    
    def _find_section(self, text: str, keywords: List[str]) -> str:
        """Find a section in the text based on keywords"""
        lines = text.split('\n')
        section_start = -1
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in keywords):
                section_start = i
                break
        
        if section_start == -1:
            return ""
        
        # Extract until next section or end
        section_lines = []
        for i in range(section_start + 1, len(lines)):
            line = lines[i]
            # Check if we hit another section header
            if self._is_section_header(line):
                break
            section_lines.append(line)
        
        return '\n'.join(section_lines)
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is likely a section header"""
        line_lower = line.lower().strip()
        headers = [
            'experience', 'education', 'skills', 'summary', 'objective',
            'certifications', 'awards', 'projects', 'publications'
        ]
        return any(header in line_lower for header in headers) and len(line.strip()) < 50
